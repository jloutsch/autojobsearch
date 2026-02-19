"""Extract profile tags from a resume file (PDF/DOC/DOCX) using Ollama."""

import io
import json
import logging
import os
import string
from urllib.parse import urlparse

import requests

logger = logging.getLogger(__name__)

_raw_ollama_url = os.environ.get("OLLAMA_URL", "http://localhost:11434")
_parsed = urlparse(_raw_ollama_url)
_allowed_hosts = {"localhost", "127.0.0.1", "host.docker.internal"}
if _parsed.scheme not in ("http", "https") or _parsed.hostname not in _allowed_hosts:
    logger.warning(
        f"OLLAMA_URL '{_raw_ollama_url}' is not a trusted local address — "
        f"falling back to http://localhost:11434"
    )
    OLLAMA_URL = "http://localhost:11434"
else:
    OLLAMA_URL = _raw_ollama_url.rstrip("/")

OLLAMA_MODEL = os.environ.get("OLLAMA_MODEL", "llama3.2:latest")

MAX_RESUME_TEXT = 8000

RESUME_PARSE_PROMPT = (
    "Analyze this resume and extract structured profile data for a job search system.\n"
    "Return ONLY valid JSON with these exact fields:\n"
    "\n"
    '- "role_tags": list of 4-8 job search keywords based on the person\'s experience '
    "and target roles. These are used as search queries on job boards. Use lowercase. "
    'Examples: "customer success", "technical account manager", "solutions engineer", '
    '"application support".\n'
    "\n"
    '- "industry_tags": list of 5-15 industry or technology area keywords that match '
    "this person's background. Used to match job descriptions. Use lowercase. "
    'Examples: "cybersecurity", "saas", "healthcare IT", "cloud infrastructure", "devops".\n'
    "\n"
    '- "skills": list of 10-30 skills extracted directly from the resume. Include ALL '
    "tools, technologies, platforms, methodologies, and competencies mentioned. "
    "Use lowercase. Include both technical skills (specific tools and platforms) and "
    "soft skills (competencies and domain expertise). "
    'Examples: "salesforce", "jira", "python", "sql", "enterprise", "onboarding", '
    '"cross-functional", "stakeholder management", "api integration", "data analysis".\n'
    "\n"
    '- "primary_role_tags": list of 2-4 ideal job titles this person is most qualified for. '
    "Use title case. These are exact titles to match against job listings for highest-priority "
    'scoring. Examples: "Customer Success Manager", "Technical Account Manager".\n'
    "\n"
    '- "secondary_role_tags": list of 3-6 secondary/stretch job titles the person could also '
    'fit. Use title case. Examples: "Solutions Engineer", "Implementation Manager", '
    '"Enterprise Support Manager".\n'
    "\n"
    '- "resume_summary": a concise 3-5 sentence professional summary of this person, written '
    "in third person, highlighting years of experience, key domains, leadership scope, and "
    "standout accomplishments. This summary is sent to an AI model alongside job descriptions "
    "to evaluate fit.\n"
    "\n"
    "IMPORTANT: For the skills field, be thorough — extract every skill, tool, platform, "
    "and technology mentioned anywhere in the resume, including the skills section, "
    "work experience bullet points, and certifications.\n"
    "\n"
    "RESUME:\n"
    "$resume_text"
)


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract plain text from a resume file based on extension."""
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        return _extract_pdf(file_bytes)
    elif ext == ".docx":
        return _extract_docx(file_bytes)
    elif ext == ".doc":
        # Try reading as .docx (works for some modern .doc files)
        try:
            return _extract_docx(file_bytes)
        except Exception:
            raise ValueError(
                "Legacy .doc format is not supported. "
                "Please save your resume as .docx or .pdf and try again."
            )
    else:
        raise ValueError(f"Unsupported file type: {ext}. Please upload a PDF or DOCX file.")


def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using pdfplumber."""
    import pdfplumber

    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    full_text = "\n\n".join(text_parts)
    if not full_text.strip():
        raise ValueError(
            "Could not extract text from the PDF. "
            "The file may be image-based or corrupt."
        )
    return full_text


def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    import docx

    doc = docx.Document(io.BytesIO(file_bytes))
    text_parts = [para.text for para in doc.paragraphs if para.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    full_text = "\n".join(text_parts)
    if not full_text.strip():
        raise ValueError(
            "Could not extract text from the document. The file may be empty or corrupt."
        )
    return full_text


def parse_resume_text(
    resume_text: str,
    progress_callback: callable = None,
) -> dict:
    """Generate profile tags from plain resume text via Ollama.

    Args:
        resume_text: Plain text of the resume (pasted by user).
        progress_callback: Optional callable(message: str) for progress updates.

    Returns:
        dict with keys: role_tags, industry_tags, skills,
        primary_role_tags, secondary_role_tags, resume_summary
    """
    if not resume_text.strip():
        raise ValueError("Resume text is empty. Paste your resume first.")

    word_count = len(resume_text.split())
    if progress_callback:
        progress_callback(f"Analyzing {word_count} words of resume text...")

    return _call_ollama(resume_text, progress_callback)


def parse_resume(
    file_bytes: bytes,
    filename: str,
    progress_callback: callable = None,
) -> dict:
    """Extract text from a resume file and generate profile tags via Ollama.

    Args:
        file_bytes: Raw file bytes.
        filename: Original filename (used to detect format).
        progress_callback: Optional callable(message: str) for progress updates.

    Returns:
        dict with keys: role_tags, industry_tags, skills,
        primary_role_tags, secondary_role_tags, resume_summary
    """
    if progress_callback:
        progress_callback("Extracting text from resume...")

    resume_text = extract_text(file_bytes, filename)

    word_count = len(resume_text.split())
    if progress_callback:
        progress_callback(f"Extracted {word_count} words from resume")

    return _call_ollama(resume_text, progress_callback)


def _call_ollama(resume_text: str, progress_callback: callable = None) -> dict:
    """Send resume text to Ollama and return structured profile tags."""
    if progress_callback:
        progress_callback("Sending resume to AI for analysis...")

    tmpl = string.Template(RESUME_PARSE_PROMPT)
    prompt = tmpl.safe_substitute(resume_text=resume_text[:MAX_RESUME_TEXT])

    try:
        resp = requests.post(
            f"{OLLAMA_URL}/api/chat",
            json={
                "model": OLLAMA_MODEL,
                "messages": [{"role": "user", "content": prompt}],
                "stream": True,
                "format": "json",
            },
            timeout=180,
            stream=True,
        )
        resp.raise_for_status()
    except requests.exceptions.ReadTimeout:
        raise ValueError(
            "AI analysis timed out. The model may be overloaded. Try again."
        )
    except requests.exceptions.ConnectionError:
        raise ValueError(
            "Could not connect to Ollama. Is it running?"
        )

    # Stream tokens and send periodic progress updates
    # Typical resume analysis response is ~400 tokens
    ESTIMATED_TOKENS = 600
    chunks = []
    token_count = 0
    last_update = 0
    try:
        for line in resp.iter_lines():
            if not line:
                continue
            try:
                chunk = json.loads(line)
            except json.JSONDecodeError:
                continue
            content = chunk.get("message", {}).get("content", "")
            if content:
                chunks.append(content)
                token_count += 1
                if progress_callback and token_count - last_update >= 15:
                    last_update = token_count
                    pct = min(95, int(token_count / ESTIMATED_TOKENS * 100))
                    progress_callback(f"AI analyzing resume... {pct}%")
            if chunk.get("done"):
                break
    except requests.exceptions.ChunkedEncodingError:
        # Partial response — use what we have if enough was generated
        if not chunks:
            raise ValueError(
                "Lost connection to AI model during analysis. Try again."
            )
        logger.warning("Ollama stream interrupted, using partial response")

    if progress_callback:
        progress_callback("Parsing AI response...")

    text = "".join(chunks).strip()

    # Handle possible markdown code block wrapping
    if text.startswith("```"):
        text = text.split("\n", 1)[1].rsplit("```", 1)[0].strip()

    result = json.loads(text)

    # Validate and sanitize all fields
    def _ensure_string_list(val, max_items=20):
        if not isinstance(val, list):
            return []
        return [
            str(item).strip()
            for item in val
            if isinstance(item, str) and item.strip()
        ][:max_items]

    parsed = {
        "role_tags": _ensure_string_list(result.get("role_tags"), max_items=10),
        "industry_tags": _ensure_string_list(result.get("industry_tags"), max_items=15),
        "skills": _ensure_string_list(result.get("skills"), max_items=30),
        "primary_role_tags": _ensure_string_list(
            result.get("primary_role_tags"), max_items=5
        ),
        "secondary_role_tags": _ensure_string_list(
            result.get("secondary_role_tags"), max_items=8
        ),
        "resume_summary": str(result.get("resume_summary", "")).strip()[:2000],
    }

    if progress_callback:
        tag_count = sum(
            len(parsed[k]) for k in ["role_tags", "industry_tags", "skills"]
        )
        progress_callback(f"Extracted {tag_count} tags and a resume summary")

    return parsed
